import logging
import os
import uuid
import json
from pathlib import Path
from typing import Optional
from urllib.parse import quote
import asyncio

from fastapi import (
    BackgroundTasks,
    APIRouter,
    Depends,
    File,
    Form,
    HTTPException,
    Request,
    UploadFile,
    status,
    Query,
)

from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session
from open_webui.internal.db import get_session, SessionLocal

from open_webui.constants import ERROR_MESSAGES
from open_webui.retrieval.vector.factory import VECTOR_DB_CLIENT

from open_webui.models.channels import Channels
from open_webui.models.users import Users
from open_webui.models.files import (
    FileForm,
    FileModel,
    FileModelResponse,
    Files,
)
from open_webui.models.chats import Chats
from open_webui.models.knowledge import Knowledges
from open_webui.models.groups import Groups
from open_webui.models.access_grants import AccessGrants


from open_webui.routers.retrieval import ProcessFileForm, process_file
from open_webui.routers.audio import transcribe

from open_webui.storage.provider import Storage


from open_webui.config import BYPASS_ADMIN_ACCESS_CONTROL
from open_webui.utils.auth import get_admin_user, get_verified_user
from open_webui.utils.misc import strict_match_mime_type
from pydantic import BaseModel

log = logging.getLogger(__name__)

router = APIRouter()


from open_webui.utils.access_control.files import has_access_to_file

############################
# Upload File
############################


def _is_text_file(file_path: str, chunk_size: int = 8192) -> bool:
    """Check if a file is likely a text file by reading a chunk and validating UTF-8.

    This catches files whose extensions are mis-mapped by mimetypes/browsers
    (e.g. TypeScript .ts → video/mp2t) without maintaining an extension whitelist.
    """
    try:
        resolved = Storage.get_file(file_path)
        with open(resolved, "rb") as f:
            chunk = f.read(chunk_size)
        if not chunk:
            return False
        # Null bytes are a strong indicator of binary content
        if b"\x00" in chunk:
            return False
        chunk.decode("utf-8")
        return True
    except (UnicodeDecodeError, Exception):
        return False


def process_uploaded_file(
    request,
    file,
    file_path,
    file_item,
    file_metadata,
    user,
    db: Optional[Session] = None,
):
    def _process_handler(db_session):
        try:
            content_type = file.content_type

            # Detect mis-labeled text files (e.g. .ts → video/mp2t)
            if content_type and content_type.startswith(("image/", "video/")):
                if _is_text_file(file_path):
                    content_type = "text/plain"

            if content_type:
                stt_supported_content_types = getattr(
                    request.app.state.config, "STT_SUPPORTED_CONTENT_TYPES", []
                )

                if strict_match_mime_type(stt_supported_content_types, content_type):
                    file_path_processed = Storage.get_file(file_path)
                    result = transcribe(
                        request, file_path_processed, file_metadata, user
                    )

                    process_file(
                        request,
                        ProcessFileForm(
                            file_id=file_item.id, content=result.get("text", "")
                        ),
                        user=user,
                        db=db_session,
                    )
                elif (not content_type.startswith(("image/", "video/"))) or (
                    request.app.state.config.CONTENT_EXTRACTION_ENGINE == "external"
                ):
                    process_file(
                        request,
                        ProcessFileForm(file_id=file_item.id),
                        user=user,
                        db=db_session,
                    )
                else:
                    raise Exception(
                        f"File type {content_type} is not supported for processing"
                    )
            else:
                log.info(
                    f"File type {file.content_type} is not provided, but trying to process anyway"
                )
                process_file(
                    request,
                    ProcessFileForm(file_id=file_item.id),
                    user=user,
                    db=db_session,
                )

        except Exception as e:
            log.error(f"Error processing file: {file_item.id}")
            Files.update_file_data_by_id(
                file_item.id,
                {
                    "status": "failed",
                    "error": str(e.detail) if hasattr(e, "detail") else str(e),
                },
                db=db_session,
            )

    if db:
        _process_handler(db)
    else:
        with SessionLocal() as db_session:
            _process_handler(db_session)


@router.post("/", response_model=FileModelResponse)
def upload_file(
    request: Request,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    metadata: Optional[dict | str] = Form(None),
    process: bool = Query(True),
    process_in_background: bool = Query(True),
    user=Depends(get_verified_user),
    db: Session = Depends(get_session),
):
    return upload_file_handler(
        request,
        file=file,
        metadata=metadata,
        process=process,
        process_in_background=process_in_background,
        user=user,
        background_tasks=background_tasks,
        db=db,
    )


def upload_file_handler(
    request: Request,
    file: UploadFile = File(...),
    metadata: Optional[dict | str] = Form(None),
    process: bool = Query(True),
    process_in_background: bool = Query(True),
    user=Depends(get_verified_user),
    background_tasks: Optional[BackgroundTasks] = None,
    db: Optional[Session] = None,
):
    log.info(f"file.content_type: {file.content_type} {process}")

    if isinstance(metadata, str):
        try:
            metadata = json.loads(metadata)
        except json.JSONDecodeError:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=ERROR_MESSAGES.DEFAULT("Invalid metadata format"),
            )
    file_metadata = metadata if metadata else {}

    try:
        unsanitized_filename = file.filename
        filename = os.path.basename(unsanitized_filename)

        file_extension = os.path.splitext(filename)[1]
        # Remove the leading dot from the file extension
        file_extension = file_extension[1:] if file_extension else ""

        if process and request.app.state.config.ALLOWED_FILE_EXTENSIONS:
            request.app.state.config.ALLOWED_FILE_EXTENSIONS = [
                ext for ext in request.app.state.config.ALLOWED_FILE_EXTENSIONS if ext
            ]

            if file_extension not in request.app.state.config.ALLOWED_FILE_EXTENSIONS:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=ERROR_MESSAGES.DEFAULT(
                        f"File type {file_extension} is not allowed"
                    ),
                )

        # replace filename with uuid
        id = str(uuid.uuid4())
        name = filename
        filename = f"{id}_{filename}"
        contents, file_path = Storage.upload_file(
            file.file,
            filename,
            {
                "OpenWebUI-User-Email": user.email,
                "OpenWebUI-User-Id": user.id,
                "OpenWebUI-User-Name": user.name,
                "OpenWebUI-File-Id": id,
            },
        )

        file_item = Files.insert_new_file(
            user.id,
            FileForm(
                **{
                    "id": id,
                    "filename": name,
                    "path": file_path,
                    "data": {
                        **({"status": "pending"} if process else {}),
                    },
                    "meta": {
                        "name": name,
                        "content_type": (
                            file.content_type
                            if isinstance(file.content_type, str)
                            else None
                        ),
                        "size": len(contents),
                        "data": file_metadata,
                    },
                }
            ),
            db=db,
        )

        if "channel_id" in file_metadata:
            channel = Channels.get_channel_by_id_and_user_id(
                file_metadata["channel_id"], user.id, db=db
            )
            if channel:
                Channels.add_file_to_channel_by_id(
                    channel.id, file_item.id, user.id, db=db
                )

        if process:
            if background_tasks and process_in_background:
                background_tasks.add_task(
                    process_uploaded_file,
                    request,
                    file,
                    file_path,
                    file_item,
                    file_metadata,
                    user,
                )
                return {"status": True, **file_item.model_dump()}
            else:
                process_uploaded_file(
                    request,
                    file,
                    file_path,
                    file_item,
                    file_metadata,
                    user,
                    db=db,
                )
                return {"status": True, **file_item.model_dump()}
        else:
            if file_item:
                return file_item
            else:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=ERROR_MESSAGES.DEFAULT("Error uploading file"),
                )

    except HTTPException as e:
        raise e
    except Exception as e:
        log.exception(e)
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.DEFAULT("Error uploading file"),
        )


############################
# List Files
############################


@router.get("/", response_model=list[FileModelResponse])
async def list_files(
    user=Depends(get_verified_user),
    content: bool = Query(True),
    db: Session = Depends(get_session),
):
    if user.role == "admin" and BYPASS_ADMIN_ACCESS_CONTROL:
        files = Files.get_files(db=db)
    else:
        files = Files.get_files_by_user_id(user.id, db=db)

    if not content:
        for file in files:
            if "content" in file.data:
                del file.data["content"]

    return files


############################
# Search Files
############################


@router.get("/search", response_model=list[FileModelResponse])
async def search_files(
    filename: str = Query(
        ...,
        description="Filename pattern to search for. Supports wildcards such as '*.txt'",
    ),
    content: bool = Query(True),
    skip: int = Query(0, ge=0, description="Number of files to skip"),
    limit: int = Query(
        100, ge=1, le=1000, description="Maximum number of files to return"
    ),
    user=Depends(get_verified_user),
    db: Session = Depends(get_session),
):
    """
    Search for files by filename with support for wildcard patterns.
    Uses SQL-based filtering with pagination for better performance.
    """
    # Determine user_id: null for admin with bypass (search all), user.id otherwise
    user_id = (
        None if (user.role == "admin" and BYPASS_ADMIN_ACCESS_CONTROL) else user.id
    )

    # Use optimized database query with pagination
    files = Files.search_files(
        user_id=user_id,
        filename=filename,
        skip=skip,
        limit=limit,
        db=db,
    )

    if not files:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No files found matching the pattern.",
        )

    if not content:
        for file in files:
            if file.data and "content" in file.data:
                del file.data["content"]

    return files


############################
# Delete All Files
############################


@router.delete("/all")
async def delete_all_files(
    user=Depends(get_admin_user), db: Session = Depends(get_session)
):
    result = Files.delete_all_files(db=db)
    if result:
        try:
            Storage.delete_all_files()
            VECTOR_DB_CLIENT.reset()
        except Exception as e:
            log.exception(e)
            log.error("Error deleting files")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=ERROR_MESSAGES.DEFAULT("Error deleting files"),
            )
        return {"message": "All files deleted successfully"}
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.DEFAULT("Error deleting files"),
        )


############################
# Get File By Id
############################


@router.get("/{id}", response_model=Optional[FileModel])
async def get_file_by_id(
    id: str, user=Depends(get_verified_user), db: Session = Depends(get_session)
):
    file = Files.get_file_by_id(id, db=db)

    if not file:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )

    if (
        file.user_id == user.id
        or user.role == "admin"
        or has_access_to_file(id, "read", user, db=db)
    ):
        return file
    else:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )


@router.get("/{id}/process/status")
async def get_file_process_status(
    id: str,
    stream: bool = Query(False),
    user=Depends(get_verified_user),
    db: Session = Depends(get_session),
):
    file = Files.get_file_by_id(id, db=db)

    if not file:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )

    if (
        file.user_id == user.id
        or user.role == "admin"
        or has_access_to_file(id, "read", user, db=db)
    ):
        if stream:
            MAX_FILE_PROCESSING_DURATION = 3600 * 2

            async def event_stream(file_id):
                # NOTE: We intentionally do NOT capture the request's db session here.
                # Each poll creates its own short-lived session to avoid holding a
                # connection for hours. A WebSocket push would be more efficient.
                for _ in range(MAX_FILE_PROCESSING_DURATION):
                    file_item = Files.get_file_by_id(file_id)  # Creates own session
                    if file_item:
                        data = file_item.model_dump().get("data", {})
                        status = data.get("status")

                        if status:
                            event = {"status": status}
                            if status == "failed":
                                event["error"] = data.get("error")

                            yield f"data: {json.dumps(event)}\n\n"
                            if status in ("completed", "failed"):
                                break
                        else:
                            # Legacy
                            break
                    else:
                        yield f"data: {json.dumps({'status': 'not_found'})}\n\n"
                        break

                    await asyncio.sleep(1)

            return StreamingResponse(
                event_stream(file.id),
                media_type="text/event-stream",
            )
        else:
            return {"status": file.data.get("status", "pending")}
    else:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )


############################
# Get File Data Content By Id
############################


@router.get("/{id}/data/content")
async def get_file_data_content_by_id(
    id: str, user=Depends(get_verified_user), db: Session = Depends(get_session)
):
    file = Files.get_file_by_id(id, db=db)

    if not file:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )

    if (
        file.user_id == user.id
        or user.role == "admin"
        or has_access_to_file(id, "read", user, db=db)
    ):
        return {"content": file.data.get("content", "")}
    else:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )


############################
# Update File Data Content By Id
############################


class ContentForm(BaseModel):
    content: str


@router.post("/{id}/data/content/update")
def update_file_data_content_by_id(
    request: Request,
    id: str,
    form_data: ContentForm,
    user=Depends(get_verified_user),
    db: Session = Depends(get_session),
):
    file = Files.get_file_by_id(id, db=db)

    if not file:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )

    if (
        file.user_id == user.id
        or user.role == "admin"
        or has_access_to_file(id, "write", user, db=db)
    ):
        try:
            process_file(
                request,
                ProcessFileForm(file_id=id, content=form_data.content),
                user=user,
                db=db,
            )
            file = Files.get_file_by_id(id=id, db=db)
        except Exception as e:
            log.exception(e)
            log.error(f"Error processing file: {file.id}")

        # Propagate content change to all knowledge collections referencing
        # this file.  Without this the old embeddings remain in the knowledge
        # collection and RAG returns both stale and current data (#20558).
        knowledges = Knowledges.get_knowledges_by_file_id(id, db=db)
        for knowledge in knowledges:
            try:
                # Remove old embeddings for this file from the KB collection
                VECTOR_DB_CLIENT.delete(
                    collection_name=knowledge.id, filter={"file_id": id}
                )
                # Re-add from the now-updated file-{file_id} collection
                process_file(
                    request,
                    ProcessFileForm(file_id=id, collection_name=knowledge.id),
                    user=user,
                    db=db,
                )
            except Exception as e:
                log.warning(
                    f"Failed to update knowledge {knowledge.id} after "
                    f"content change for file {id}: {e}"
                )

        return {"content": file.data.get("content", "")}
    else:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )


############################
# Get File Content By Id
############################


@router.get("/{id}/content")
async def get_file_content_by_id(
    id: str,
    user=Depends(get_verified_user),
    attachment: bool = Query(False),
    db: Session = Depends(get_session),
):
    file = Files.get_file_by_id(id, db=db)

    if not file:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )

    if (
        file.user_id == user.id
        or user.role == "admin"
        or has_access_to_file(id, "read", user, db=db)
    ):
        try:
            file_path = Storage.get_file(file.path)
            file_path = Path(file_path)

            # Check if the file already exists in the cache
            if file_path.is_file():
                # Handle Unicode filenames
                filename = file.meta.get("name", file.filename)
                encoded_filename = quote(filename)  # RFC5987 encoding

                content_type = file.meta.get("content_type")
                filename = file.meta.get("name", file.filename)
                encoded_filename = quote(filename)
                headers = {}

                if attachment:
                    headers["Content-Disposition"] = (
                        f"attachment; filename*=UTF-8''{encoded_filename}"
                    )
                else:
                    if content_type == "application/pdf" or filename.lower().endswith(
                        ".pdf"
                    ):
                        headers["Content-Disposition"] = (
                            f"inline; filename*=UTF-8''{encoded_filename}"
                        )
                        content_type = "application/pdf"
                    elif content_type != "text/plain":
                        headers["Content-Disposition"] = (
                            f"attachment; filename*=UTF-8''{encoded_filename}"
                        )

                return FileResponse(file_path, headers=headers, media_type=content_type)

            else:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail=ERROR_MESSAGES.NOT_FOUND,
                )
        except HTTPException as e:
            raise e
        except Exception as e:
            log.exception(e)
            log.error("Error getting file content")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=ERROR_MESSAGES.DEFAULT("Error getting file content"),
            )
    else:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )


@router.get("/{id}/content/html")
async def get_html_file_content_by_id(
    id: str, user=Depends(get_verified_user), db: Session = Depends(get_session)
):
    file = Files.get_file_by_id(id, db=db)

    if not file:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )

    file_user = Users.get_user_by_id(file.user_id, db=db)
    if not file_user.role == "admin":
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )

    if (
        file.user_id == user.id
        or user.role == "admin"
        or has_access_to_file(id, "read", user, db=db)
    ):
        try:
            file_path = Storage.get_file(file.path)
            file_path = Path(file_path)

            # Check if the file already exists in the cache
            if file_path.is_file():
                log.info(f"file_path: {file_path}")
                return FileResponse(file_path)
            else:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail=ERROR_MESSAGES.NOT_FOUND,
                )
        except HTTPException as e:
            raise e
        except Exception as e:
            log.exception(e)
            log.error("Error getting file content")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=ERROR_MESSAGES.DEFAULT("Error getting file content"),
            )
    else:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )



############################
# Get File Content in Format (txt/docx/hwpx)
############################


@router.get("/{id}/{format}/content")
async def get_file_content_in_format(
    id: str,
    format: str,
    user=Depends(get_verified_user),
    attachment: bool = Query(True),
    db: Session = Depends(get_session),
):
    """Serve file content converted to txt/docx/hwpx format."""
    if format not in ("txt", "docx", "hwpx"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Unsupported format. Use txt, docx, or hwpx.",
        )

    file = Files.get_file_by_id(id, db=db)
    if not file:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )

    if not (
        file.user_id == user.id
        or user.role == "admin"
        or has_access_to_file(id, "read", user, db=db)
    ):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )

    # Get text content from file data (transcription, extraction, etc.)
    content = ""
    if file.data and isinstance(file.data, dict):
        content = file.data.get("content", "")

    if not content:
        # Fallback: try to read original file as text
        try:
            file_path = Storage.get_file(file.path)
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        except Exception:
            content = ""

    if not content:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="No text content available for this file.",
        )

    from io import BytesIO

    base_name = Path(file.meta.get("name", file.filename)).stem

    if format == "txt":
        filename = f"{base_name}.txt"
        encoded = quote(filename)
        headers = {}
        if attachment:
            headers["Content-Disposition"] = (
                f"attachment; filename*=UTF-8\'\'{encoded}"
            )
        return StreamingResponse(
            BytesIO(content.encode("utf-8")),
            media_type="text/plain; charset=utf-8",
            headers=headers,
        )

    elif format == "docx":
        from docx import Document as DocxDocument

        doc = DocxDocument()
        for para in content.split("\n"):
            doc.add_paragraph(para)
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        filename = f"{base_name}.docx"
        encoded = quote(filename)
        headers = {}
        if attachment:
            headers["Content-Disposition"] = (
                f"attachment; filename*=UTF-8\'\'{encoded}"
            )
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers,
        )

    elif format == "hwpx":
        import httpx
        import re as _re

        mcp_url = "http://220.124.155.35:5002/mcp"
        try:
            async with httpx.AsyncClient(timeout=120) as client:
                common_h = {
                    "Content-Type": "application/json",
                    "Accept": "application/json, text/event-stream",
                }
                # 1) MCP initialize
                init_resp = await client.post(mcp_url, json={
                    "jsonrpc": "2.0", "method": "initialize", "id": 1,
                    "params": {
                        "protocolVersion": "2025-06-18",
                        "capabilities": {},
                        "clientInfo": {"name": "owui-file-route", "version": "1.0"},
                    },
                }, headers=common_h)
                session_id = init_resp.headers.get("mcp-session-id", "")
                sess_h = {**common_h, "Mcp-Session-Id": session_id}

                # 2) initialized notification
                await client.post(mcp_url, json={
                    "jsonrpc": "2.0", "method": "notifications/initialized",
                }, headers=sess_h)

                # 3) generate_hwp
                call_resp = await client.post(mcp_url, json={
                    "jsonrpc": "2.0", "method": "tools/call", "id": 2,
                    "params": {
                        "name": "generate_hwp",
                        "arguments": {
                            "content": content,
                            "file_name": base_name,
                            "user_id": "system",
                            "template_type": "default",
                            "doc_title": base_name,
                        },
                    },
                }, headers=sess_h)

                # 4) Parse SSE or JSON response for download URL
                resp_text = call_resp.text
                # Handle 202 Accepted (SSE stream)
                if call_resp.status_code == 202:
                    async with client.stream("GET", mcp_url, headers=sess_h, timeout=120) as stream:
                        async for chunk in stream.aiter_text():
                            resp_text += chunk
                            if '"result"' in resp_text or '"error"' in resp_text:
                                break

                # Extract URL from markdown link [name](url) or plain URL
                url_match = _re.search(r'\(https?://[^)\s]+\.hwpx?[^)\s]*)\)', resp_text)
                if not url_match:
                    url_match = _re.search(r'(https?://[^\s"]+\.hwpx?[^\s"]*)', resp_text)

                if url_match:
                    file_url = url_match.group(1)
                    # Download the generated hwpx file
                    dl_resp = await client.get(file_url, timeout=60)
                    if dl_resp.status_code == 200:
                        filename = f"{base_name}.hwpx"
                        encoded = quote(filename)
                        headers = {}
                        if attachment:
                            headers["Content-Disposition"] = (
                                f"attachment; filename*=UTF-8\'\'{encoded}"
                            )
                        return StreamingResponse(
                            BytesIO(dl_resp.content),
                            media_type="application/hwpx",
                            headers=headers,
                        )

                # Cleanup
                try:
                    await client.delete(mcp_url, headers={"Mcp-Session-Id": session_id})
                except Exception:
                    pass

        except Exception as e:
            log.warning(f"MCP hwpx generation failed: {e}")

        # Fallback: serve as plain text if MCP fails
        filename = f"{base_name}.txt"
        encoded = quote(filename)
        headers = {}
        if attachment:
            headers["Content-Disposition"] = (
                f"attachment; filename*=UTF-8\'\'{encoded}"
            )
        return StreamingResponse(
            BytesIO(content.encode("utf-8")),
            media_type="text/plain; charset=utf-8",
            headers=headers,
        )

############################
# Dynamic HWPX Generation (양식 기반 문서 자동생성)
############################


class HwpxGenerateForm(BaseModel):
    template_file_id: str  # 양식 HWPX 파일 ID
    content_text: Optional[str] = None  # 직접 입력한 내용 텍스트
    content_file_id: Optional[str] = None  # 내용 소스 파일 ID (PDF 등)
    model: Optional[str] = None  # 사용할 LLM 모델 (None이면 기본 task 모델)
    doc_title: Optional[str] = None  # 출력 파일명


@router.post("/generate-hwpx")
async def generate_hwpx_dynamic_endpoint(
    request: Request,
    form_data: HwpxGenerateForm,
    user=Depends(get_verified_user),
    db: Session = Depends(get_session),
):
    """
    양식 HWPX + 내용 텍스트/파일 → AI 분석 → HWPX 문서 자동생성

    1. 양식 파일에서 경량 XML 추출
    2. 내용 소스 텍스트 확보
    3. AI에게 양식 + 내용 전달 → 명령 JSON 수신
    4. 명령 실행하여 HWPX 생성
    """
    # ── TEMP DEBUG: 1차 구조분석 + 마커 분리까지만 실행 (2a/2b 건너뜀) ──
    # role 분류 알고리즘 개선 작업 중. 되돌리려면 이 return 블록 삭제.
    return await debug_hwpx_structure_endpoint(request, form_data, user, db)

    from io import BytesIO
    from open_webui.utils.hwpx_analyzer import (
        analyze_hwpx,
        truncate_xml,
        build_structure_analysis_prompt,
        parse_structure_from_llm,
        build_chapter_classify_prompt,
        parse_chapter_classify_from_llm,
        build_section_fill_prompt,
        parse_section_fill_from_llm,
        pdf_to_base64_images,
        pdf_to_text,
    )
    from open_webui.utils.hwp_generator import assemble_hwpx_hybrid
    from open_webui.utils.chat import generate_chat_completion
    from open_webui.utils.task import get_task_model_id

    # 1) 양식 파일 가져오기
    template_file = Files.get_file_by_id(form_data.template_file_id, db=db)
    if not template_file:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="양식 파일을 찾을 수 없습니다",
        )

    template_path = Storage.get_file(template_file.path)
    if not template_path:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="양식 파일 경로를 찾을 수 없습니다",
        )

    # 2) 내용 확보 (PDF → 텍스트 + JPEG 이미지)
    content_text = form_data.content_text or ""
    content_images = None
    pdf_text_content = ""

    if form_data.content_file_id:
        content_file = Files.get_file_by_id(form_data.content_file_id, db=db)
        if content_file:
            content_type = content_file.meta.get("content_type", "")
            file_name = content_file.meta.get("name", content_file.filename)

            if content_type == "application/pdf" or file_name.lower().endswith(".pdf"):
                content_path = Storage.get_file(content_file.path)

                # PDF → 텍스트 추출 (2차 호출에 사용)
                try:
                    pdf_text_content = pdf_to_text(content_path)
                    log.info(f"PDF 텍스트 추출 완료: {len(pdf_text_content):,}자")
                except Exception as e:
                    log.warning(f"PDF 텍스트 추출 실패: {e}")

                # PDF → JPEG 이미지 변환 (텍스트와 병행)
                try:
                    content_images = pdf_to_base64_images(content_path)
                    log.info(f"PDF → JPEG 변환 완료: {len(content_images)}페이지")
                except Exception as e:
                    log.warning(f"PDF 이미지 변환 실패: {e}")
                    content_images = None

            # PDF가 아닌 경우 → 텍스트 추출본 사용
            if not pdf_text_content and content_images is None and not content_text:
                content_text = content_file.data.get("content", "") if content_file.data else ""

    if not content_text and not content_images and not pdf_text_content:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="작성할 내용이 없습니다. content_text 또는 content_file_id를 제공하세요",
        )

    # 3) 양식 분석 (경량 XML 추출 + 축소)
    try:
        analysis = analyze_hwpx(template_path)
        light_xml = analysis["light_xml"]
        log.info(
            f"양식 분석 완료: 문단 {analysis['paragraph_count']}개, "
            f"표 {analysis['table_count']}개, "
            f"경량 XML {len(light_xml):,}B"
        )
        truncate_result = truncate_xml(light_xml)
        truncated_xml = truncate_result["xml"]
        removed_indices = truncate_result["removed_indices"]
        idx_map = truncate_result.get("idx_map")
        log.info(f"XML 축소: {len(light_xml):,} → {len(truncated_xml):,}자, 제거 문단 {len(removed_indices)}개")
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"양식 분석 실패: {e}",
        )

    # 4) 모델 확인
    models = request.app.state.MODELS
    model_id = form_data.model
    if not model_id:
        model_id = get_task_model_id(
            "",
            request.app.state.config.TASK_MODEL,
            request.app.state.config.TASK_MODEL_EXTERNAL,
            models,
        )

    if not model_id or model_id not in models:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"사용 가능한 모델이 없습니다: {model_id}",
        )

    # --- AI 호출 헬퍼 ---
    async def _call_llm(messages, task_name):
        payload = {
            "model": model_id,
            "messages": messages,
            "stream": False,
            "metadata": {"task": task_name},
        }
        resp = await generate_chat_completion(
            request, form_data=payload, user=user
        )
        if hasattr(resp, "body"):
            resp = json.loads(resp.body.decode("utf-8"))
        if "error" in resp:
            raise HTTPException(
                status_code=status.HTTP_502_BAD_GATEWAY,
                detail=f"AI 응답 오류 ({task_name}): {resp['error']}",
            )
        return resp["choices"][0]["message"]["content"]

    # 5a) 1차 AI 호출 — 양식 구조 분석 (role + description + marker + table)
    try:
        messages_1 = build_structure_analysis_prompt(truncated_xml, auto_truncate=False)
        for i, m in enumerate(messages_1):
            content_preview = m.get("content", "")
            if isinstance(content_preview, str):
                content_preview = content_preview[:2000]
            log.info(f"[HWP-DEBUG] 1차 요청 messages[{i}] role={m.get('role')}, content(앞2000자):\n{content_preview}")

        llm_content_1 = await _call_llm(messages_1, "hwpx_structure_analysis")
        log.info(f"[HWP-DEBUG] 1차 LLM 응답 (전체 {len(llm_content_1)}자):\n{llm_content_1}")

        structure = parse_structure_from_llm(llm_content_1)

        log.info(
            f"1차 구조 분석 완료: 문단 {len(structure.get('paragraphs', []))}개, "
            f"표 {len(structure.get('tables', []))}개 (level은 아직 없음)"
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"양식 구조 분석 실패: {e}",
        )

    # 5a-2) 1b/1c/1d/1e 파이프라인
    try:
        from open_webui.utils.hwpx_analyzer import (
            build_level_analysis_prompt,
            parse_level_from_llm,
            merge_levels_into_structure,
            build_role_classification_prompt,
            parse_role_classification_from_llm,
            merge_roles_into_structure,
            build_chapter_types_from_structure,
            compute_parent_instance_children,
            build_exclusivity_analysis_prompt,
            parse_exclusivity_from_llm,
            compute_format_observations,
            build_format_analysis_prompt,
            parse_format_rules_from_llm,
        )

        # 1b: level
        messages_level = build_level_analysis_prompt(structure, signals=None)
        log.info(f"[HWP-DEBUG] 1b 요청: {len(structure.get('paragraphs', []))}개 문단 level 판단")
        llm_content_level = await _call_llm(messages_level, "hwpx_1b_level")
        level_map = parse_level_from_llm(llm_content_level)
        log.info(f"[HWP-DEBUG] 1b 파싱: {len(level_map)}개 level 결정")
        structure = merge_levels_into_structure(structure, level_map)

        # 1c: role 분류
        messages_role = build_role_classification_prompt(structure, signals=None)
        log.info(f"[HWP-DEBUG] 1c 요청: role 분류")
        llm_content_role = await _call_llm(messages_role, "hwpx_1c_role")
        role_map = parse_role_classification_from_llm(llm_content_role)
        log.info(f"[HWP-DEBUG] 1c 파싱: {len(role_map)}개 role 결정")
        structure = merge_roles_into_structure(structure, role_map)

        # 1d: 배타 규칙
        pc_data = compute_parent_instance_children(structure)
        exclusive_rules = []
        if pc_data:
            role_markers = {}
            for p in structure.get("paragraphs", []):
                r, m = p.get("role", ""), p.get("marker", "")
                if r and m and r not in role_markers:
                    role_markers[r] = m
            messages_excl = build_exclusivity_analysis_prompt(pc_data, role_markers=role_markers)
            try:
                llm_content_excl = await _call_llm(messages_excl, "hwpx_1d_exclusivity")
                exclusive_rules = parse_exclusivity_from_llm(llm_content_excl)
                log.info(f"[HWP-DEBUG] 1d 파싱: 배타 규칙 {len(exclusive_rules)}개")
            except Exception as e:
                log.warning(f"[HWP-DEBUG] 1d 배타 규칙 판정 실패: {e}")
                exclusive_rules = []
            if exclusive_rules:
                structure["exclusive_rules"] = exclusive_rules

        # 1e: format / blank 규칙
        format_obs = compute_format_observations(structure, light_xml, idx_map=idx_map)
        if format_obs.get("role_formats") or format_obs.get("transitions"):
            messages_format = build_format_analysis_prompt(format_obs)
            try:
                llm_content_format = await _call_llm(messages_format, "hwpx_1e_format")
                parsed_format = parse_format_rules_from_llm(llm_content_format)
                fmt_rules = parsed_format.get("format_rules", {})
                blnk_rules = parsed_format.get("blank_rules", [])
                log.info(
                    f"[HWP-DEBUG] 1e 파싱: format_rules {len(fmt_rules)}개, "
                    f"blank_rules {len(blnk_rules)}개"
                )
                if fmt_rules:
                    structure["format_rules"] = fmt_rules
                if blnk_rules:
                    structure["blank_rules"] = blnk_rules
            except Exception as e:
                log.warning(f"[HWP-DEBUG] 1e 판정 실패: {e}")

        # chapter_types 생성 (level 기반 트리 + variant 분리)
        structure = build_chapter_types_from_structure(structure)

        log.info(f"[HWP-DEBUG] 1.5차 후 최종 structure:\n{json.dumps(structure, ensure_ascii=False, indent=2)}")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"level 분석 실패: {e}",
        )

    # ── 5b) 2a AI 호출 — 소스 대제목 추출 + 양식 타입 분류 ──
    chapter_types = structure.get("chapter_types", {})
    if not chapter_types:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="chapter_types 생성 실패. 양식에 level 1 문단이 없을 수 있습니다.",
        )

    # header role 목록 추출 — 속성 기반:
    # level 0 + 첫 level 1 문단 이전에 등장 + role 이름 무관 (동적 양식 대응)
    first_ch_idx = next(
        (p.get("idx", 0) for p in structure.get("paragraphs", []) if p.get("level", 0) == 1),
        float("inf"),
    )
    header_roles = []
    for p in structure.get("paragraphs", []):
        role = p.get("role", "")
        if not role or role in header_roles:
            continue
        if p.get("level", 0) != 0:
            continue
        if p.get("idx", 0) >= first_ch_idx:
            continue
        header_roles.append(role)

    try:
        messages_2a = build_chapter_classify_prompt(
            chapter_types,
            header_roles,
            content_text=content_text,
            content_images=content_images,
            pdf_text=pdf_text_content,
        )
        log.info(f"[HWP-DEBUG] 2a 요청: chapter_types={list(chapter_types.keys())}, header_roles={header_roles}")

        llm_content_2a = await _call_llm(messages_2a, "hwpx_chapter_classify")
        log.info(f"[HWP-DEBUG] 2a LLM 응답 (전체 {len(llm_content_2a)}자):\n{llm_content_2a}")

        classify_result = parse_chapter_classify_from_llm(llm_content_2a)
        chapters = classify_result.get("chapters", [])
        header_data = classify_result.get("header", {})

        log.info(f"[HWP-DEBUG] 2a 파싱: {len(chapters)}개 대제목, header={list(header_data.keys())}")
        log.info(f"[HWP-DEBUG] 2a chapters:\n{json.dumps(chapters, ensure_ascii=False, indent=2)}")

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"소스 대제목 분류 실패: {e}",
        )

    # ── 5c) 2b AI 호출 — 섹션별 콘텐츠 채우기 (대제목마다 1회) ──
    # 소스 텍스트를 대제목별로 분할
    from open_webui.utils.hwpx_analyzer import (
        _extract_texts_by_idx,
        split_source_by_chapters,
    )

    chapter_titles_list = [ch.get("title", "") for ch in chapters]
    if pdf_text_content:
        source_sections = split_source_by_chapters(pdf_text_content, chapter_titles_list)
    else:
        source_sections = [""] * len(chapters)

    # role 카탈로그 구성 (1차 structure에서 추출)
    idx_texts = {}
    if truncated_xml:
        try:
            idx_texts = _extract_texts_by_idx(truncated_xml)
        except Exception:
            pass

    full_role_catalog = {}
    for p in structure.get("paragraphs", []):
        role = p.get("role", "")
        if role and role not in full_role_catalog:
            sample = idx_texts.get(p.get("idx", -1), "")
            full_role_catalog[role] = {
                "description": p.get("description", ""),
                "marker": p.get("marker", ""),
                "level": p.get("level", 0),
                "sample": sample,
            }

    body_items = []
    for ch_idx, chapter in enumerate(chapters):
        ch_type = chapter.get("type", "")
        ch_title = chapter.get("title", "")

        if ch_type not in chapter_types:
            log.warning(f"[HWP-DEBUG] 2b 스킵: 알 수 없는 타입 '{ch_type}' (대제목: {ch_title})")
            continue

        type_info = chapter_types[ch_type]
        pattern = type_info.get("pattern", {})
        title_role = type_info.get("title_role", "chapter_title")

        # 이 패턴에 사용되는 role만 카탈로그에서 추출
        def _collect_roles(pat: dict) -> set:
            roles = set()
            for rname, info in pat.items():
                roles.add(rname)
                children = info.get("children", {})
                if children:
                    roles.update(_collect_roles(children))
            return roles

        pattern_roles = _collect_roles(pattern)
        section_catalog = {r: full_role_catalog[r] for r in pattern_roles if r in full_role_catalog}

        # 이 섹션에 해당하는 소스 텍스트만 전달
        section_pdf_text = source_sections[ch_idx] if ch_idx < len(source_sections) else ""

        try:
            messages_2b = build_section_fill_prompt(
                ch_title,
                ch_type,
                pattern,
                section_catalog,
                content_text=content_text,
                content_images=content_images,
                pdf_text=section_pdf_text,
                exclusive_rules=structure.get("exclusive_rules", []),
                format_rules=structure.get("format_rules", {}),
            )

            log.info(f"[HWP-DEBUG] 2b[{ch_idx}] 요청: type={ch_type}, title={ch_title}, roles={list(section_catalog.keys())}")

            llm_content_2b = await _call_llm(messages_2b, f"hwpx_section_fill_{ch_idx}")
            log.info(f"[HWP-DEBUG] 2b[{ch_idx}] LLM 응답 ({len(llm_content_2b)}자):\n{llm_content_2b}")

            items = parse_section_fill_from_llm(llm_content_2b)
            log.info(f"[HWP-DEBUG] 2b[{ch_idx}] 파싱: {len(items)}개 항목")

            # 대제목 항목 추가 (2b 출력에는 대제목 자체가 없으므로)
            body_items.append({"role": title_role, "text": ch_title})
            body_items.extend(items)

        except Exception as e:
            log.warning(f"2b[{ch_idx}] 실패 ({ch_title}): {e}")
            # 실패한 섹션은 대제목만 추가
            body_items.append({"role": title_role, "text": ch_title})

    # ── 6) 결과 조합 → assemble_hwpx_hybrid ──
    content = {
        "header": header_data,
        "body": body_items,
    }
    log.info(
        f"[HWP-DEBUG] 최종 콘텐츠: header={list(header_data.keys())}, "
        f"body={len(body_items)}개 항목"
    )
    log.info(f"[HWP-DEBUG] body 내용:\n{json.dumps(body_items, ensure_ascii=False, indent=2)}")

    # 7) HWPX 조립
    try:
        result = assemble_hwpx_hybrid(
            template_path,
            structure=structure,
            content=content,
            removed_indices=removed_indices,
            idx_map=truncate_result.get("idx_map"),
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"HWPX 생성 실패: {e}",
        )

    # 7-1) 결과 문서 내용 디버그 덤프
    try:
        from hwpx import HwpxDocument
        debug_doc = HwpxDocument.open(BytesIO(result.data))
        debug_lines = []
        debug_lines.append(f"=== 결과 문서: 문단 {len(debug_doc.paragraphs)}개 ===")
        for i, p in enumerate(debug_doc.paragraphs):
            text = p.text or ""
            preview = text[:100].replace("\n", "\\n")
            if len(text) > 100:
                preview += "..."
            debug_lines.append(f"  P[{i}] {preview}")
        tables_found = []
        for p in debug_doc.paragraphs:
            tables_found.extend(p.tables)
        debug_lines.append(f"=== 결과 문서: 표 {len(tables_found)}개 ===")
        for t_idx, tbl in enumerate(tables_found):
            rows = int(tbl.element.get("rowCnt", "0"))
            cols = int(tbl.element.get("colCnt", "0"))
            debug_lines.append(f"  T[{t_idx}] {rows}x{cols}")
            NS = "{http://www.hancom.co.kr/hwpml/2011/paragraph}"
            for tr in tbl.element.findall(f"{NS}tr"):
                for tc in tr.findall(f"{NS}tc"):
                    addr = tc.find(f"{NS}cellAddr")
                    r = addr.get("rowAddr", "?") if addr is not None else "?"
                    c = addr.get("colAddr", "?") if addr is not None else "?"
                    cell_text = ""
                    for t_elem in tc.iter(f"{NS}t"):
                        if t_elem.text:
                            cell_text += t_elem.text
                    preview = cell_text[:80].replace("\n", "\\n")
                    if len(cell_text) > 80:
                        preview += "..."
                    if cell_text.strip():
                        debug_lines.append(f"    ({r},{c}) {preview}")
        log.info(f"[HWP-DEBUG] 결과 문서 내용:\n" + "\n".join(debug_lines))
    except Exception as e:
        log.warning(f"[HWP-DEBUG] 결과 문서 덤프 실패: {e}")

    # 7) 응답
    doc_title = form_data.doc_title or template_file.meta.get("name", "document")
    if not doc_title.endswith(".hwpx"):
        doc_title = f"{Path(doc_title).stem}.hwpx"

    encoded = quote(doc_title)
    headers = {
        "Content-Disposition": f"attachment; filename*=UTF-8''{encoded}",
        "X-Hwpx-Success": str(result.success_count),
        "X-Hwpx-Fail": str(result.fail_count),
    }
    if result.errors:
        from urllib.parse import quote as url_quote
        headers["X-Hwpx-Errors"] = url_quote("; ".join(result.errors))

    return StreamingResponse(
        BytesIO(result.data),
        media_type="application/hwpx",
        headers=headers,
    )


############################
# Debug: HWPX 1차 구조분석 + 마커 분리까지만 (2a/2b 건너뜀)
############################


@router.post("/debug-hwpx-structure")
async def debug_hwpx_structure_endpoint(
    request: Request,
    form_data: HwpxGenerateForm,
    user=Depends(get_verified_user),
    db: Session = Depends(get_session),
):
    """
    디버그: 1차 AI 구조분석 + _split_roles_by_marker 까지만 실행하고 JSON으로 반환.
    2a/2b 는 건너뜀. 결과는 /tmp/hwpx_debug_last.json 에도 덤프.
    """
    import copy
    import re
    from open_webui.utils.hwpx_analyzer import (
        analyze_hwpx,
        truncate_xml,
        build_structure_analysis_prompt,
        _split_roles_by_marker,
        _normalize_marker_type,
        _repair_json,
        pdf_to_text,
    )
    from open_webui.utils.chat import generate_chat_completion
    from open_webui.utils.task import get_task_model_id

    # 1) 양식 파일
    template_file = Files.get_file_by_id(form_data.template_file_id, db=db)
    if not template_file:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="양식 파일을 찾을 수 없습니다",
        )
    template_path = Storage.get_file(template_file.path)
    if not template_path:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="양식 파일 경로를 찾을 수 없습니다",
        )

    # 2) 소스(선택) — 1차엔 쓰지 않지만 컨텍스트 기록용
    content_text = form_data.content_text or ""
    pdf_text_content = ""
    if form_data.content_file_id:
        content_file = Files.get_file_by_id(form_data.content_file_id, db=db)
        if content_file:
            content_type = content_file.meta.get("content_type", "")
            file_name = content_file.meta.get("name", content_file.filename)
            if content_type == "application/pdf" or file_name.lower().endswith(".pdf"):
                content_path = Storage.get_file(content_file.path)
                try:
                    pdf_text_content = pdf_to_text(content_path)
                except Exception as e:
                    log.warning(f"[DEBUG-HWPX] PDF 텍스트 추출 실패: {e}")

    # 3) 양식 분석
    try:
        analysis = analyze_hwpx(template_path)
        light_xml = analysis["light_xml"]
        truncate_result = truncate_xml(light_xml)
        truncated_xml = truncate_result["xml"]
        removed_indices = truncate_result["removed_indices"]
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"양식 분석 실패: {e}",
        )

    # 4) 모델
    models = request.app.state.MODELS
    model_id = form_data.model
    if not model_id:
        model_id = get_task_model_id(
            "",
            request.app.state.config.TASK_MODEL,
            request.app.state.config.TASK_MODEL_EXTERNAL,
            models,
        )
    if not model_id or model_id not in models:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"사용 가능한 모델이 없습니다: {model_id}",
        )

    # 5) 1차 AI 호출
    messages_1 = build_structure_analysis_prompt(truncated_xml, auto_truncate=False)
    payload = {
        "model": model_id,
        "messages": messages_1,
        "stream": False,
        "metadata": {"task": "hwpx_structure_analysis_debug"},
    }
    try:
        resp = await generate_chat_completion(request, form_data=payload, user=user)
        if hasattr(resp, "body"):
            resp = json.loads(resp.body.decode("utf-8"))
        if "error" in resp:
            raise HTTPException(
                status_code=status.HTTP_502_BAD_GATEWAY,
                detail=f"AI 응답 오류: {resp['error']}",
            )
        llm_raw = resp["choices"][0]["message"]["content"]
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"1차 AI 호출 실패: {e}",
        )

    # 6) 파싱 (split 전 상태 보존 위해 인라인)
    json_match = re.search(r'```(?:json)?\s*([\[{][\s\S]*?[\]}])\s*```', llm_raw)
    if json_match:
        raw = json_match.group(1)
    else:
        brace_match = re.search(r'\{[\s\S]*\}', llm_raw)
        if brace_match:
            raw = brace_match.group(0)
        else:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="1차 응답에서 JSON을 찾을 수 없음",
            )
    try:
        data = json.loads(raw, strict=False)
    except json.JSONDecodeError:
        try:
            data = json.loads(_repair_json(raw), strict=False)
        except json.JSONDecodeError as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"1차 JSON 파싱 실패: {e}",
            )
    if not isinstance(data, dict) or "paragraphs" not in data:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="1차 응답에 'paragraphs' 키 없음",
        )

    paragraphs_before = copy.deepcopy(data.get("paragraphs", []))

    # 7) 마커 분리 적용
    paragraphs_after = _split_roles_by_marker(copy.deepcopy(paragraphs_before))

    # 8) split_log — idx 기준 before/after 비교
    before_by_idx = {p.get("idx"): p.get("role", "") for p in paragraphs_before}
    split_log = []
    for p in paragraphs_after:
        idx = p.get("idx")
        after_role = p.get("role", "")
        before_role = before_by_idx.get(idx, "")
        if before_role != after_role:
            split_log.append({
                "idx": idx,
                "marker": p.get("marker", ""),
                "marker_type": _normalize_marker_type(p.get("marker", "")),
                "before_role": before_role,
                "after_role": after_role,
            })

    # 9) marker_normalization 맵 (role → {marker_type → [markers]})
    marker_norm = {}
    for p in paragraphs_before:
        role = p.get("role", "")
        marker = p.get("marker", "")
        if not role:
            continue
        mt = _normalize_marker_type(marker)
        marker_norm.setdefault(role, {}).setdefault(mt, set()).add(marker)
    marker_norm_serializable = {
        role: {mt: sorted(list(markers)) for mt, markers in by_type.items()}
        for role, by_type in marker_norm.items()
    }

    # 10) 응답 구성
    result_payload = {
        "success": True,
        "template": {
            "file_id": form_data.template_file_id,
            "name": template_file.meta.get("name", template_file.filename),
            "paragraph_count": analysis.get("paragraph_count", 0),
            "table_count": analysis.get("table_count", 0),
        },
        "source": {
            "file_id": form_data.content_file_id,
            "text_length": len(pdf_text_content or content_text),
        },
        "xml": {
            "light_xml_size": len(light_xml),
            "truncated_xml_size": len(truncated_xml),
            "removed_indices_count": len(removed_indices),
            "removed_indices": removed_indices,
        },
        "model_id": model_id,
        "stage_1_structure_analysis": {
            "prompt_messages": messages_1,
            "llm_raw_response": llm_raw,
            "structure_before_split": {
                "paragraphs": paragraphs_before,
                "tables": data.get("tables", []),
            },
            "structure_after_split": {
                "paragraphs": paragraphs_after,
                "tables": data.get("tables", []),
            },
            "split_log": split_log,
            "marker_normalization": marker_norm_serializable,
        },
    }

    # 11) 디스크 덤프 (/tmp/hwpx_debug_last.json)
    dump_path = "/tmp/hwpx_debug_last.json"
    try:
        with open(dump_path, "w", encoding="utf-8") as f:
            json.dump(result_payload, f, ensure_ascii=False, indent=2)
        result_payload["debug_file_path"] = dump_path
        log.info(f"[DEBUG-HWPX] 덤프 완료: {dump_path}")
    except Exception as e:
        log.warning(f"[DEBUG-HWPX] 덤프 실패: {e}")
        result_payload["debug_file_path"] = None

    return result_payload


############################
# Analyze HWPX Template (양식 분석 미리보기)
############################


@router.post("/analyze-hwpx/{file_id}")
async def analyze_hwpx_endpoint(
    file_id: str,
    user=Depends(get_verified_user),
    db: Session = Depends(get_session),
):
    """양식 HWPX 파일을 분석하여 경량 XML과 메타정보를 반환합니다."""
    from open_webui.utils.hwpx_analyzer import analyze_hwpx

    file = Files.get_file_by_id(file_id, db=db)
    if not file:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="파일을 찾을 수 없습니다",
        )

    file_path = Storage.get_file(file.path)
    if not file_path:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="파일 경로를 찾을 수 없습니다",
        )

    try:
        result = analyze_hwpx(file_path)
        return {
            "file_id": file_id,
            "filename": file.meta.get("name", file.filename),
            "paragraph_count": result["paragraph_count"],
            "table_count": result["table_count"],
            "light_xml_size": len(result["light_xml"]),
            "original_xml_size": len(result["original_xml"]),
            "reduction_pct": round(
                (1 - len(result["light_xml"]) / len(result["original_xml"])) * 100
            ),
        }
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"양식 분석 실패: {e}",
        )


############################
# Dynamic HWPX Generation v2 (역할 기반 문서 생성)
############################


class HwpxGenerateV2Form(BaseModel):
    template_file_id: str
    content_text: Optional[str] = None
    content_file_id: Optional[str] = None
    model: Optional[str] = None
    doc_title: Optional[str] = None


@router.post("/generate-hwpx-v2")
async def generate_hwpx_v2_endpoint(
    request: Request,
    form_data: HwpxGenerateV2Form,
    user=Depends(get_verified_user),
    db: Session = Depends(get_session),
):
    """
    역할 기반 HWPX 문서 생성 v2

    1. 양식에서 서식 그룹 추출 (코드)
    2. 1차 AI: 서식 그룹 → 역할 해석
    3. 2차 AI: 역할 + 소스 → 콘텐츠 생성
    4. 역할 기반 문서 조립
    """
    from io import BytesIO
    from open_webui.utils.hwpx_analyzer import (
        extract_style_groups,
        build_role_interpret_prompt,
        parse_role_interpret_from_llm,
        build_role_content_prompt,
        parse_role_content_from_llm,
        pdf_to_base64_images,
        pdf_to_text,
    )
    from open_webui.utils.hwp_generator import assemble_hwpx
    from open_webui.utils.chat import generate_chat_completion
    from open_webui.utils.task import get_task_model_id

    # 1) 양식 파일
    template_file = Files.get_file_by_id(form_data.template_file_id, db=db)
    if not template_file:
        raise HTTPException(status_code=404, detail="양식 파일을 찾을 수 없습니다")
    template_path = Storage.get_file(template_file.path)
    if not template_path:
        raise HTTPException(status_code=404, detail="양식 파일 경로를 찾을 수 없습니다")

    # 2) 내용 확보
    content_text = form_data.content_text or ""
    content_images = None
    pdf_text_content = ""

    if form_data.content_file_id:
        content_file = Files.get_file_by_id(form_data.content_file_id, db=db)
        if content_file:
            content_type = content_file.meta.get("content_type", "")
            file_name = content_file.meta.get("name", content_file.filename)
            if content_type == "application/pdf" or file_name.lower().endswith(".pdf"):
                content_path = Storage.get_file(content_file.path)
                try:
                    pdf_text_content = pdf_to_text(content_path)
                except Exception as e:
                    log.warning(f"PDF 텍스트 추출 실패: {e}")
                try:
                    content_images = pdf_to_base64_images(content_path)
                except Exception as e:
                    log.warning(f"PDF 이미지 변환 실패: {e}")
            if not pdf_text_content and content_images is None and not content_text:
                content_text = content_file.data.get("content", "") if content_file.data else ""

    if not content_text and not content_images and not pdf_text_content:
        raise HTTPException(status_code=400, detail="작성할 내용이 없습니다")

    # 3) 서식 그룹 추출 (코드)
    try:
        style_catalog = extract_style_groups(template_path)
        log.info(f"서식 그룹 {len(style_catalog['groups'])}개 추출")
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"서식 그룹 추출 실패: {e}")

    # 4) 모델 확인
    models = request.app.state.MODELS
    model_id = form_data.model
    if not model_id:
        model_id = get_task_model_id(
            "", request.app.state.config.TASK_MODEL,
            request.app.state.config.TASK_MODEL_EXTERNAL, models,
        )
    if not model_id or model_id not in models:
        raise HTTPException(status_code=400, detail=f"모델 없음: {model_id}")

    async def _call_llm(messages, task_name):
        payload = {
            "model": model_id, "messages": messages,
            "stream": False, "metadata": {"task": task_name},
        }
        resp = await generate_chat_completion(request, form_data=payload, user=user)
        if hasattr(resp, "body"):
            resp = json.loads(resp.body.decode("utf-8"))
        if "error" in resp:
            raise HTTPException(status_code=502, detail=f"AI 오류 ({task_name}): {resp['error']}")
        return resp["choices"][0]["message"]["content"]

    # 5) 1차 AI: 역할 해석
    try:
        messages_1 = build_role_interpret_prompt(style_catalog)
        llm_1 = await _call_llm(messages_1, "hwpx_role_interpret")
        role_map = parse_role_interpret_from_llm(llm_1)
        log.info(f"역할 해석 완료: {len(role_map)}개")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"역할 해석 실패: {e}")

    # 6) 2차 AI: 콘텐츠 생성
    try:
        messages_2 = build_role_content_prompt(
            style_catalog, role_map,
            content_text=content_text,
            content_images=content_images,
            pdf_text=pdf_text_content,
        )
        llm_2 = await _call_llm(messages_2, "hwpx_role_content")
        content_data = parse_role_content_from_llm(llm_2)
        log.info(f"콘텐츠 생성 완료: body {len(content_data.get('body', []))}개 항목")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"콘텐츠 생성 실패: {e}")

    # 7) 문서 조립
    try:
        result = assemble_hwpx(template_path, style_catalog, role_map, content_data)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"문서 조립 실패: {e}")

    # 8) 응답
    doc_title = form_data.doc_title or template_file.meta.get("name", "document")
    if not doc_title.endswith(".hwpx"):
        doc_title = f"{Path(doc_title).stem}.hwpx"

    encoded = quote(doc_title)
    headers = {
        "Content-Disposition": f"attachment; filename*=UTF-8''{encoded}",
        "X-Hwpx-Success": str(result.success_count),
        "X-Hwpx-Fail": str(result.fail_count),
    }
    if result.errors:
        from urllib.parse import quote as url_quote
        headers["X-Hwpx-Errors"] = url_quote("; ".join(result.errors[:10]))

    return StreamingResponse(
        BytesIO(result.data),
        media_type="application/hwpx",
        headers=headers,
    )


@router.get("/{id}/content/{file_name}")
async def get_file_content_by_id(
    id: str, user=Depends(get_verified_user), db: Session = Depends(get_session)
):
    file = Files.get_file_by_id(id, db=db)

    if not file:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )

    if (
        file.user_id == user.id
        or user.role == "admin"
        or has_access_to_file(id, "read", user, db=db)
    ):
        file_path = file.path

        # Handle Unicode filenames
        filename = file.meta.get("name", file.filename)
        encoded_filename = quote(filename)  # RFC5987 encoding
        headers = {
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
        }

        if file_path:
            file_path = Storage.get_file(file_path)
            file_path = Path(file_path)

            # Check if the file already exists in the cache
            if file_path.is_file():
                return FileResponse(file_path, headers=headers)
            else:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail=ERROR_MESSAGES.NOT_FOUND,
                )
        else:
            # File path doesn’t exist, return the content as .txt if possible
            file_content = file.content.get("content", "")
            file_name = file.filename

            # Create a generator that encodes the file content
            def generator():
                yield file_content.encode("utf-8")

            return StreamingResponse(
                generator(),
                media_type="text/plain",
                headers=headers,
            )
    else:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )


############################
# Delete File By Id
############################


@router.delete("/{id}")
async def delete_file_by_id(
    id: str, user=Depends(get_verified_user), db: Session = Depends(get_session)
):
    file = Files.get_file_by_id(id, db=db)

    if not file:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )

    if (
        file.user_id == user.id
        or user.role == "admin"
        or has_access_to_file(id, "write", user, db=db)
    ):

        # Clean up KB associations and embeddings before deleting
        knowledges = Knowledges.get_knowledges_by_file_id(id, db=db)
        for knowledge in knowledges:
            # Remove KB-file relationship
            Knowledges.remove_file_from_knowledge_by_id(knowledge.id, id, db=db)
            # Clean KB embeddings (same logic as /knowledge/{id}/file/remove)
            try:
                VECTOR_DB_CLIENT.delete(
                    collection_name=knowledge.id, filter={"file_id": id}
                )
                if file.hash:
                    VECTOR_DB_CLIENT.delete(
                        collection_name=knowledge.id, filter={"hash": file.hash}
                    )
            except Exception as e:
                log.debug(f"KB embedding cleanup for {knowledge.id}: {e}")

        result = Files.delete_file_by_id(id, db=db)
        if result:
            try:
                Storage.delete_file(file.path)
                VECTOR_DB_CLIENT.delete(collection_name=f"file-{id}")
            except Exception as e:
                log.exception(e)
                log.error("Error deleting files")
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=ERROR_MESSAGES.DEFAULT("Error deleting files"),
                )
            return {"message": "File deleted successfully"}
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=ERROR_MESSAGES.DEFAULT("Error deleting file"),
            )
    else:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=ERROR_MESSAGES.NOT_FOUND,
        )
